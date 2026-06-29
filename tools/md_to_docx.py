#!/usr/bin/env python3
"""
Markdown -> Word (.docx) converter using python-docx.
Handles: headings H1-H4, bold/italic inline, bullet lists, tables, blockquotes,
         horizontal rules, code spans, and paragraph text.
Usage: python tools/md_to_docx.py <input.md> <output.docx>
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def apply_inline(run_parent, text):
    """Parse **bold**, *italic*, `code` and plain text into runs."""
    # Pattern: **bold**, *italic* (but not **), `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            run_parent.add_run(text[pos:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            r = run_parent.add_run(m.group(2))
            r.bold = True
        elif full.startswith('*'):
            r = run_parent.add_run(m.group(3))
            r.italic = True
        elif full.startswith('`'):
            r = run_parent.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        pos = m.end()
    if pos < len(text):
        run_parent.add_run(text[pos:])

def add_styled_para(doc, text, style='Normal', indent=0):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.4)
    apply_inline(p, text)
    return p

def parse_table_rows(lines):
    """Collect contiguous | lines as a table block."""
    rows = []
    for ln in lines:
        s = ln.strip()
        if not s.startswith('|'):
            break
        # skip separator rows like |---|---|
        if re.match(r'^\|[\s\-:|\s]+\|$', s):
            continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        rows.append(cells)
    return rows

def render_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            apply_inline(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacing after table

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        raw_lines = f.readlines()

    doc = Document()

    # Global styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set heading styles
    for level, size in [(1, 18), (2, 15), (3, 13), (4, 11)]:
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        if level == 1:
            hstyle.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            hstyle.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            hstyle.font.color.rgb = RGBColor(0x5A, 0x96, 0xC8)

    lines = [ln.rstrip('\n') for ln in raw_lines]
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # Fenced code block
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                # render collected code
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    r = p.add_run(cl)
                    r.font.name = 'Courier New'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                code_lines = []
                i += 1
                continue
        if in_code_block:
            code_lines.append(ln)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            # strip trailing # if any
            text = re.sub(r'\s+#+\s*$', '', text)
            p = doc.add_heading(level=level)
            p.clear()
            apply_inline(p, text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Table block — collect all consecutive table lines
        if stripped.startswith('|'):
            # gather run
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j])
                j += 1
            rows = parse_table_rows(table_lines)
            render_table(doc, rows)
            i = j
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = re.sub(r'^>\s?', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # light gray left border via shading
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '4472C4')
            pBdr.append(left)
            pPr.append(pBdr)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[-*+]\s+', stripped):
            text = re.sub(r'^[-*+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            apply_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^\d+\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            apply_inline(p, m.group(1))
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        apply_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Saved {docx_path}')

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python tools/md_to_docx.py input.md output.docx')
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
